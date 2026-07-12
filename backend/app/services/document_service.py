from pathlib import Path
import shutil
import uuid

from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.models.document import Document
from app.models.user import User
from pypdf import PdfReader
from docx import Document as DocxDocument
from app.llm.service import LLMService
from app.vectorstore.chroma import collection

llm_service = LLMService()

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

ALLOWED_TYPES = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
def extract_pdf_text(file_path: Path) -> str:
    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def extract_docx_text(file_path: Path) -> str:
    document = DocxDocument(file_path)

    return "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )


def extract_txt_text(file_path: Path) -> str:
    return file_path.read_text(
        encoding="utf-8",
    )

def chunk_text(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 200,
) -> list[str]:
    chunks = []

    start = 0

    while start < len(text):
        end = start + chunk_size

        chunks.append(text[start:end])

        start += chunk_size - overlap

    return chunks

def save_document(
        
    db: Session,
    file: UploadFile,
    current_user: User,
):
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type",
        )

    user_dir = UPLOAD_DIR / f"user_{current_user.id}"
    user_dir.mkdir(exist_ok=True)

    extension = Path(file.filename).suffix

    unique_filename = f"{uuid.uuid4()}{extension}"

    file_path = user_dir / unique_filename

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    document = Document(
        filename=file.filename,
        filepath=str(file_path),
        content_type=file.content_type,
        user_id=current_user.id,
    )

    db.add(document)
    db.commit()
    db.refresh(document)
    if file.content_type == "application/pdf":
        text = extract_pdf_text(file_path)

    elif (
    file.content_type
    == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = extract_docx_text(file_path)

    else:
        text = extract_txt_text(file_path)

    chunks = chunk_text(text)
    
    for index, chunk in enumerate(chunks):
        embedding = llm_service.embedding(chunk)

        collection.add(
            ids=[
                f"{document.id}_{index}"
            ],
        embeddings=[embedding],
        documents=[chunk],
        metadatas=[
            {
                "document_id": document.id,
                "user_id": current_user.id,
                "chunk_index": index,
            }
        ],
    )
      

    print(f"Total Chunks: {len(chunks)}")

    for i, chunk in enumerate(chunks):
        print(f"\n------ Chunk {i+1} ------")
        print(chunk[:200])
    return document

def get_user_documents(
    db: Session,
    current_user: User,
):
    return (
        db.query(Document)
        .filter(Document.user_id == current_user.id)
        .order_by(Document.id.desc())
        .all()
    )

def delete_document(
    db: Session,
    document_id: int,
    current_user: User,
):
    document = (
        db.query(Document)
        .filter(
            Document.id == document_id,
            Document.user_id == current_user.id,
        )
        .first()
    )

    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found",
        )

    file_path = Path(document.filepath)

    if file_path.exists():
        file_path.unlink()
    
    collection.delete(
    where={
        "document_id": document.id,
    }
)

    db.delete(document)
    db.commit()

    return {
        "message": "Document deleted successfully"
    }